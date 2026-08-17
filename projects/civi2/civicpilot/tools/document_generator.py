"""DocumentGenerator wrapper (python-docx & reportlab PDF) (Requirement 15.3).

Generates application-support docx and PDF documents for a single scheme candidate.
Accepts only a single scheme's criteria/info payload -- structurally preventing
cross-scheme data leakage.
"""

from __future__ import annotations

import io
import re
from typing import Any

import docx
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Spacer

_XML_CONTROL_CHAR_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")


def _clean_xml_str(text: Any) -> str:
    return _XML_CONTROL_CHAR_RE.sub("", str(text))


class DocumentGenerator:
    """Document generator producing python-docx and ReportLab PDF formatted files."""

    def generate_document(self, scheme_payload: dict[str, Any]) -> bytes:
        """Generates a docx document for a single scheme payload."""
        doc = docx.Document()

        scheme_name = _clean_xml_str(
            scheme_payload.get("scheme_name", "Government Scheme Application Support")
        )
        doc.add_heading(scheme_name, level=1)

        summary = scheme_payload.get("summary")
        if summary:
            doc.add_heading("Scheme Overview", level=2)
            doc.add_paragraph(_clean_xml_str(summary))

        criteria = scheme_payload.get("criteria", [])
        if criteria:
            doc.add_heading("Eligibility Criteria", level=2)
            for item in criteria:
                doc.add_paragraph(_clean_xml_str(item), style="List Bullet")

        application_steps = scheme_payload.get("application_steps", [])
        if application_steps:
            doc.add_heading("Application Steps", level=2)
            for idx, step in enumerate(application_steps, start=1):
                doc.add_paragraph(f"{idx}. {_clean_xml_str(step)}")

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def generate_pdf_document(self, scheme_payload: dict[str, Any]) -> bytes:
        """Generates a valid, beautifully formatted PDF document binary."""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=54,
            leftMargin=54,
            topMargin=54,
            bottomMargin=54,
        )

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "DocTitle",
            parent=styles["Heading1"],
            fontSize=18,
            leading=22,
            textColor=colors.HexColor("#1E1B4B"),
            spaceAfter=8,
        )
        heading_style = ParagraphStyle(
            "DocHeading",
            parent=styles["Heading2"],
            fontSize=13,
            leading=16,
            textColor=colors.HexColor("#4338CA"),
            spaceBefore=12,
            spaceAfter=6,
        )
        body_style = ParagraphStyle(
            "DocBody",
            parent=styles["Normal"],
            fontSize=10,
            leading=14,
            textColor=colors.HexColor("#1F2937"),
            spaceAfter=6,
        )
        bullet_style = ParagraphStyle(
            "DocBullet",
            parent=styles["Normal"],
            fontSize=10,
            leading=14,
            leftIndent=15,
            textColor=colors.HexColor("#374151"),
            spaceAfter=4,
        )

        elements = []

        scheme_name = _clean_xml_str(
            scheme_payload.get("scheme_name", "Government Scheme Application Guide")
        )
        elements.append(Paragraph(f"🏛️ {scheme_name}", title_style))
        elements.append(Paragraph("<b>CivicPilot Official Application & Eligibility Guide</b>", body_style))
        elements.append(
            HRFlowable(
                width="100%",
                thickness=1.5,
                color=colors.HexColor("#6366F1"),
                spaceBefore=6,
                spaceAfter=12,
            )
        )

        summary = scheme_payload.get("summary")
        if summary:
            elements.append(Paragraph("Scheme Overview", heading_style))
            elements.append(Paragraph(_clean_xml_str(summary), body_style))

        criteria = scheme_payload.get("criteria", [])
        if criteria:
            elements.append(Paragraph("Key Eligibility Criteria", heading_style))
            for crit in criteria:
                elements.append(Paragraph(f"• {_clean_xml_str(crit)}", bullet_style))

        documents = scheme_payload.get("required_documents", [])
        if documents:
            elements.append(Paragraph("Required Documents Checklist", heading_style))
            for doc_item in documents:
                elements.append(Paragraph(f"✓ {_clean_xml_str(doc_item)}", bullet_style))

        steps = scheme_payload.get("application_steps", [])
        if steps:
            elements.append(Paragraph("Step-by-Step Application Process", heading_style))
            for idx, step in enumerate(steps, start=1):
                elements.append(Paragraph(f"<b>Step {idx}:</b> {_clean_xml_str(step)}", bullet_style))

        portal_url = scheme_payload.get("portal_url", "https://myscheme.gov.in")
        elements.append(Spacer(1, 12))
        elements.append(
            Paragraph(
                f"<b>Official Portal:</b> <font color='#4F46E5'><u>{_clean_xml_str(portal_url)}</u></font>",
                body_style,
            )
        )

        doc.build(elements)
        return buffer.getvalue()
